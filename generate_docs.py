import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

# --- CONFIGURATION ---
INPUT_FILE = "api_specification.md"
OUTPUT_DOCX = "Siren_Marketplace_API_Spec.docx"
OUTPUT_PDF = "Siren_Marketplace_API_Spec.pdf"

# --- HELPER FUNCTIONS ---

def read_markdown(filename):
    """Reads the markdown file."""
    if not os.path.exists(filename):
        print(f"Error: {filename} not found. Please ensure the file is in the directory.")
        return []
    with open(filename, "r", encoding="utf-8") as f:
        return f.readlines()

def clean_text(text):
    """Removes markdown bold/italic syntax for plain text rendering."""
    text = text.replace("**", "").replace("*", "").strip()
    return text

# --- DOCX GENERATOR ---

def create_docx(lines):
    doc = Document()

    # 1. Setup Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Code Block Style
    styles = doc.styles
    try:
        code_style = styles.add_style('CodeBlock', 1) # 1 is Paragraph style
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        code_style.paragraph_format.space_after = Pt(2)
        code_style.paragraph_format.left_indent = Inches(0.2)
    except:
        pass # Style might already exist

    # Parsing State
    in_code_block = False
    in_table = False
    table_data = []

    for line in lines:
        line = line.rstrip()

        # 1. Handle Code Blocks
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(line, style='CodeBlock')
            p.paragraph_format.space_after = Pt(0)
            continue

        # 2. Handle Tables
        if line.strip().startswith("|"):
            if "---" in line: # Skip separator line
                continue
            in_table = True
            # Parse row
            row_cells = [cell.strip() for cell in line.strip("|").split("|")]
            table_data.append(row_cells)
            continue
        else:
            if in_table:
                # Flush Table
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for r, row in enumerate(table_data):
                        for c, cell_text in enumerate(row):
                            if c < len(table.rows[r].cells):
                                table.rows[r].cells[c].text = clean_text(cell_text)
                                # Bold header
                                if r == 0:
                                    run = table.rows[r].cells[c].paragraphs[0].runs[0]
                                    run.font.bold = True
                table_data = []
                in_table = False
                doc.add_paragraph("") # Spacing after table

        # 3. Handle Headings
        if line.startswith("# "):
            doc.add_heading(clean_text(line[2:]), level=0)
        elif line.startswith("## "):
            doc.add_heading(clean_text(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(clean_text(line[4:]), level=2)
        elif line.startswith("#### "):
            doc.add_heading(clean_text(line[5:]), level=3)

        # 4. Handle Lists
        elif line.strip().startswith("- "):
            p = doc.add_paragraph(clean_text(line[2:]), style='List Bullet')

        # 5. Handle Blockquotes / Notes
        elif line.strip().startswith(">"):
            p = doc.add_paragraph(clean_text(line.replace(">", "")), style='Intense Quote')

        # 6. Horizontal Rules
        elif line.strip() == "---":
            p = doc.add_paragraph()
            run = p.add_run("________________________________________________________________")
            run.font.color.rgb = RGBColor(200, 200, 200)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 7. Normal Text (Skip empty lines if multiple)
        elif line.strip():
            doc.add_paragraph(clean_text(line))

    doc.save(OUTPUT_DOCX)
    print(f"✓ DOCX created: {OUTPUT_DOCX}")

# --- PDF GENERATOR ---

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 10)
        self.set_text_color(150)
        self.cell(0, 10, 'Siren Marketplace API Specification', 0, 1, 'R')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(150)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def create_pdf(lines):
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    in_code_block = False
    in_table = False
    table_data = []

    for line in lines:
        line = line.rstrip()

        # 1. Code Blocks
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if in_code_block:
                pdf.ln(2)
                pdf.set_font("Courier", "", 9)
                pdf.set_fill_color(245, 245, 245) # Light gray bg
                pdf.set_text_color(0, 0, 0)
            else:
                pdf.set_font("Arial", "", 11) # Reset font
                pdf.ln(2)
            continue

        if in_code_block:
            # Simulate code block with background
            pdf.cell(0, 5, txt=line, ln=True, border=0, fill=True)
            continue

        # 2. Tables
        if line.strip().startswith("|"):
            if "---" in line: continue
            in_table = True
            row_cells = [cell.strip() for cell in line.strip("|").split("|")]
            table_data.append(row_cells)
            continue
        else:
            if in_table:
                if table_data:
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 10)
                    # Simple table rendering
                    col_width = pdf.w / len(table_data[0]) - 10
                    row_height = 8

                    for r_idx, row in enumerate(table_data):
                        if r_idx == 0: pdf.set_font("Arial", "B", 10)
                        else: pdf.set_font("Arial", "", 10)

                        for item in row:
                            pdf.cell(col_width, row_height, txt=clean_text(item), border=1)
                        pdf.ln(row_height)
                    pdf.ln(5)
                    pdf.set_font("Arial", "", 11)

                table_data = []
                in_table = False

        # 3. Headings
        if line.startswith("# "):
            pdf.ln(10)
            pdf.set_font("Arial", "B", 20)
            pdf.set_text_color(44, 62, 80) # Dark Blue
            pdf.cell(0, 10, clean_text(line[2:]), ln=True)
            pdf.ln(5)
            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(0, 0, 0)

        elif line.startswith("## "):
            pdf.ln(8)
            pdf.set_font("Arial", "B", 16)
            pdf.set_text_color(44, 62, 80)
            pdf.cell(0, 10, clean_text(line[3:]), ln=True)
            pdf.ln(2)
            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(0, 0, 0)

        elif line.startswith("### "):
            pdf.ln(5)
            pdf.set_font("Arial", "B", 13)
            pdf.set_text_color(52, 73, 94)
            pdf.cell(0, 10, clean_text(line[4:]), ln=True)
            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(0, 0, 0)

        # 4. Blockquotes
        elif line.strip().startswith(">"):
            pdf.ln(2)
            pdf.set_font("Arial", "I", 11)
            pdf.set_text_color(80, 80, 80)
            pdf.multi_cell(0, 6, clean_text(line.replace(">", "")))
            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(0, 0, 0)

        # 5. Normal Text
        elif line.strip():
            pdf.multi_cell(0, 6, clean_text(line))

        # 6. Spacing
        else:
            pdf.ln(2)

    pdf.output(OUTPUT_PDF)
    print(f"✓ PDF created: {OUTPUT_PDF}")

# --- MAIN EXECUTION ---

if __name__ == "__main__":
    content = read_markdown(INPUT_FILE)
    if content:
        print("Generating documents...")
        create_docx(content)
        create_pdf(content)
        print("Done! Check your folder.")